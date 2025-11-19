import os
import shutil
import io
import uuid
import re
import base64
import traceback
import logging
from typing import IO, Callable
import xml.etree.ElementTree as ET
from pathlib import Path

import tiktoken
import tiktoken.load
from PIL import Image
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document

logger = logging.getLogger(__name__)


class WordSpliter:

    def __init__(
            self,
            docx_object: str | IO[bytes],
            filepath: str | None = None,
            max_tokens: int = 256,
            image_identificator: Callable | None = None,
            encoder_path: str | None = None,
            is_merge: bool = False,
        ):

        if isinstance(docx_object, bytes):
            if filepath is None:
                raise ValueError("filepath must be given while docx is bytes")
            self.doc_name = filepath
        else:
            self.doc_name =Path(docx_object).stem
        self.max_tokens = max_tokens

        self.doc = Document(docx_object)
        self.image_identificator = image_identificator
        self._encoder_path = encoder_path
        self.is_merge: bool = is_merge

        # exchange & cache
        self.output_doc_dir: str | None = None
        self._father_file = ""
        self._grandpa_file = ""
        self._current_filename: str | None = None
        self._last_level = 0
        self._current_first_title_level = 0
        self._current_second_title_level = 0
        self._current_third_title_level = 0
        self._first_router = ""
        self._second_router = ""
        self._frame = []
        self._num_and_title: str = ''
        self.current_output_dir: str = ''

        # statistic
        self._block_count = 0
        self._heading_count = 0
        self._table_count = 0
        self._image_count = 0

        self.FLOWCHART_NS = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'v': 'urn:schemas-microsoft-com:vml',
            'o': 'urn:schemas-microsoft-com:office:office',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
        }

    def _set_output_path(self, path: str):
        """set output path"""
        self.output_doc_dir = os.path.join(path, self.doc_name)
        self.current_output_dir = self.output_doc_dir
        os.makedirs(self.output_doc_dir, exist_ok=True)

    @property
    def _above_text(self):
        """current content"""
        return self._grandpa_file + self._father_file

    def iter_block_items(self, parent):
        """Generator to yield each paragraph and table in document order"""
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        else:
            raise ValueError("Invalid parent element")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
            else:
                logger.debug(f"skip: {type(child)} is a unknown child type")
                yield child

    def is_heading(self, para):
        """
        is H1/H2/H3
        """
        style_name = para.style.name.lower() if para.style else "no style"
        if style_name == "title":
            return True
        return (style_name.startswith("heading")
                and style_name[-1].isdigit()
                and int(style_name[-1]) <= 3)

    def is_toc(self, para):
        """
        it toc
        """
        style_name = para.style.name.lower() if para.style else "no style"
        return style_name.startswith("toc")

    def has_auto_num(self, para):
        """
        has auto number indexing
        """
        p_xml = para._element.xml
        root = ET.fromstring(p_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        num_pr = root.find('.//w:numPr', ns)
        return num_pr is None

    def get_heading_level(self, para):
        """get title layer 
             0: main title
             1/2/3: sub title (levels)
        """
        if not para.style:
            return -1

        style_name = para.style.name.lower()
        if style_name.startswith("heading") and style_name[-1].isdigit():
            level = int(style_name[-1])
            logger.info(f"get title<level={level}>")
            return level
        return -1

    def extract_checkboxes(self, paragraph):
        """
        get checking box informations
          checking: ☑
          unchecked: ☐
        """
        try:
            p_xml = paragraph._element.xml

            if 'w:fldChar' in p_xml and 'FORMCHECKBOX' in p_xml:
                default_elements = re.findall(r'<w:default[^>]*w:val="([^"]*)"', p_xml)
                if default_elements:
                    default_value = default_elements[0]
                    if default_value == "1":
                        return "☑"
                    else:
                        return "☐"

                # 如果没有找到默认值，尝试其他方式 -> 查找复选框的当前状态
                if 'w:checked' in p_xml:
                    return "☑"
                else:
                    return "☐"
            return None

        except Exception as e:
            logger.debug(f'get checking info error: {traceback.format_exc()}')
            logger.warning(f"'get checking info: {e}")
            return None
    
    def has_image_in_table(self, para):
        """is there images in table/para"""
        para_xml = para._element.xml
        if para._element.xpath('.//pic:pic'):
            return True
        if 'w:object' in para_xml and 'v:imagedata r:id' in para_xml:
            return True
        return False

    def process_cell(self, cell):
        """
        process one cell content to markdown format
        """
        cell_content = []

        for para in cell.paragraphs:
            checkbox_content = self.extract_checkboxes(para)
            strike_text = self.extract_strike_through_text(para)

            if para.text.strip():
                word_txt = re.sub(r'\n', ' ', para.text)
                cell_txt = word_txt.strip()

                if checkbox_content:
                    cell_txt = f"{checkbox_content} {cell_txt}"
                if strike_text:
                    cell_txt = f"~~{cell_txt}~~"
                cell_content.append(cell_txt)

            if self.has_image_in_table(para):
                logger.info("found images in cell")
                txt = self.process_image(para, True)
                cell_content.append(txt)

        if cell.tables:
            logger.info("nested table in cell")
            for nested_table in cell.tables:
                nested_content = self.process_nested_table(nested_table)
                if nested_content.strip():
                    cell_content.append(nested_content.replace("|", "\\|"))
        return cell_content
    
    def build_markdown_table(self, rows):
        """
        build markdown table from rows
        """
        header = "|" + "|".join(rows[0]) + "|"
        separator = "|" + "|".join(["---"] * len(rows[0])) + "|"
        body = "\n".join("|" + "|".join(row) + "|" for row in rows[1:])
        table_content = f"{header}\n{separator}\n{body}"

        if table_content:
            pass
            # logger.debug("调用LLM优化表格内容")  # todo 暂停使用
            # table_content = llm.table_to_text(self._above_text, table_content)
        else:
            logger.warning("生成的表格内容为空")
        return table_content

    def process_table(self, table: Table) -> str:
        """
        process table to markdown format
        """
        try:
            rows = []

            for i, row in enumerate(table.rows):
                row_cells = []
                for _, cell in enumerate(row.cells):
                    cell_content = self.process_cell(cell)
                    current_content = " ".join(filter(None, cell_content))
                    row_cells.append(current_content)

                # 处理水平合并单元格：只有当一整行全部都是连续相同的内容时，才只保留第一个
                if len(row_cells) > 1:
                    first_cell = row_cells[0]
                    all_same = all(cell == first_cell for cell in row_cells)
                    if all_same and first_cell.strip():
                        processed_row = [first_cell] + [""] * (len(row_cells) - 1)
                    else:
                        processed_row = row_cells
                else:
                    processed_row = row_cells
                rows.append(processed_row)

            col_count = len(rows[0]) if rows else 0
            for i, r in enumerate(rows):
                if len(r) != col_count:
                    logger.warning(f"row <{i}> has wrong col counts: ({len(r)} != {col_count}), auto padding")
                    rows[i] += [""] * (col_count - len(r))

            table_content = self.build_markdown_table(rows)
            return table_content
        except Exception as e:
            logger.debug(traceback.format_exc())
            logger.error(f"error occured while table processing: {e}")
            return " "

    def extract_strike_through_text(self, paragraph) -> str:
        """
        find extract strike line
        """
        try:
            strike_texts = []
            for run in paragraph.runs:
                if run.font.strike:
                    strike_texts.append(run.text)
            if strike_texts:
                return " ".join(strike_texts)
            return ""
        except Exception as e:
            logger.debug(f"error occured while finding strike lines: {traceback.format_exc()}")
            logger.error(f"error occured while finding strike lines: {e}")
            return ""

    def process_nested_table(self, table: Table) -> str:
        """process nested table to markdown format"""
        logger.debug("处理嵌套表格")
        try:
            rows = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_content = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_content.append(para.text)
                        if para._element.xpath('.//pic:pic'):
                            txt = self.process_table_image(para)
                            if txt.strip():
                                cell_content.append(txt)
                    if cell.tables:
                        for nested_table in cell.tables:
                            nested_content = self.process_nested_table(nested_table)
                            if nested_content.strip():
                                cell_content.append(nested_content)
                    row_cells.append("<br>".join(filter(None, cell_content)))
                rows.append(row_cells)
            if not rows:
                logger.warning("nested table is empty")
                return ''

            return self.build_markdown_table(rows)
        except Exception as e:
            logger.debug(f"error occured while processing nested table: {traceback.format_exc()}")
            logger.error(f"error occured while processing nested table: {e}")
            return ""

    def process_table_image(self, para):
        """process images in table"""
        try:
            pics = []
            for image in para._element.xpath('.//pic:pic'):
                for image_id in image.xpath('.//a:blip/@r:embed'):
                    image_part = self.doc.part.related_parts[image_id]
                    image_name = os.path.basename(image_part.partname)
                    logger.info(f"==>processing image:{image_name}")

                    image_bytes = image_part.blob

                    if image_name.lower().endswith('.emf'):
                        logger.debug(f"emf image found, should transform to png: {image_name}")
                        try:
                            image_bytes, image_name = self.convert_emf_to_png(image_bytes, image_name)
                            logger.info(f"EMF -> PNG: {image_name}")
                        except Exception as conv_e:
                            logger.warning(f"EMF transform fail: {str(conv_e)}, use original emf image")

                    pic_path = os.path.join(self.current_output_dir, image_name)
                    doc_pic_path = f"./{image_name}"

                    logger.info(f"save iamges (in table): {pic_path}")
                    with open(pic_path, 'wb') as f:
                        f.write(image_bytes)

                    if callable(self.image_identificator):
                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                        image_info = self.image_identificator(
                            image_base64,
                            self._above_text + "\n `there is a image in table`"
                        )
                    else:
                        image_info = "this is picture desc but fake not real generate by llm"

                    pics.append(f"![image]({doc_pic_path}) :  ({image_info})")
            return ' '.join(pics)
        except Exception as e:
            logger.debug(f"error occured while processing imaged table: {traceback.format_exc()}")
            logger.error(f"error occured while processing imaged table: {e}")
            return ""

    def convert_emf_to_png(self, emf_bytes: bytes, original_name: str) -> tuple[bytes, str]:
        """emf to png"""
        emf_stream = io.BytesIO(emf_bytes)

        with Image.open(emf_stream) as img:
            logger.debug(f"read emf: {img.format} -> {img.size}")
            if img.mode != 'RGB':
                img = img.convert('RGB')

            png_stream = io.BytesIO()
            img.save(png_stream, format='PNG')
            png_bytes = png_stream.getvalue()

            new_name = original_name.rsplit('.', 1)[0] + '.png'
            logger.info(f"emf -> png: png size {len(png_bytes)} bytes")
            return png_bytes, new_name


    def process_image(self, paragraph, is_table) -> str:
        """处理文档中的图片"""
        try:
            para_xml = paragraph._element.xml
            image_list = []
            if paragraph._element.xpath('.//pic:pic'):
                """处理的png图片"""
                for image in paragraph._element.xpath('.//pic:pic'):
                    for image_id in image.xpath('.//a:blip/@r:embed'):
                        image_part = self.doc.part.related_parts[image_id]
                        self.save_image(image_list, image_part, is_table)

            if 'v:imagedata r:id' in para_xml:
                for rel_id, part in self.doc.part.related_parts.items():
                    if (hasattr(part, 'content_type') and
                            'image' in part.content_type and
                            hasattr(part, 'blob')):
                        if f'v:imagedata r:id="{rel_id}"' in para_xml:
                            self.save_image(image_list, part, is_table)

            if image_list:
                return "  ".join(image_list)
        except Exception as e:
            logger.debug(f"processing image failed: {traceback.format_exc()}")
            logger.error(f"processing image failed: {e}")
        return ""

    def save_image(self, image_list, image_part, is_table):
        
        image_name = os.path.basename(image_part.partname)
        logger.info(f"==>处理图片:{image_name}")
        image_bytes = image_part.blob
        if image_name.lower().endswith('.emf'):
            image_bytes, image_name = self.convert_emf_to_png(image_bytes, image_name)
            if not image_name:
                logger.debug("convert emf/wmf to png failed, skip this image")
                return

        pic_path = os.path.join(self.current_output_dir, image_name)
        doc_pic_path = f"./{image_name}"
        logger.info(f"保存图片: {pic_path}")

        os.makedirs(self.current_output_dir, exist_ok=True)
        with open(pic_path, 'wb') as f:
            f.write(image_bytes)

        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        if callable(self.image_identificator):
            image_info = self.image_identification(
                image_base64, self._above_text + "\n`there is an image above`"
            )
        else:
            image_info = "this is picture desc but fake not real generate by llm"
        
        if is_table:
            image_list.append(f"![image]({doc_pic_path})  {image_info}")
        else:
            image_list.append(f"![image]({doc_pic_path})  \n{image_info}")

    def clean_title(self, txt):
        first_txt = re.sub(r'[\\/]', ' or ', txt)
        clean_txt = re.sub(r' {2}', ' ', first_txt)
        return clean_txt.replace(":", "").strip()

    def _parse_header_level_0(self, para):
        """
        Docstring for _parse_header_level_0
        
        :param self: Description
        :param para: Description
        """
        # if self._last_level == 0:
        #     self._num_and_title += " " + title_text
        # if self._last_level == -1:
        #     self._grandpa_file = self._father_file = content_txt = ""
        #     self._num_and_title = f"0 {title_text}"
        #     self._last_level = 0

        # self._current_first_title_level = 0
        # self._current_second_title_level = 0
        # self.current_output_dir = self.output_doc_dir
        return
    
    def _parse_header_level_1(self, para, level: int, title: str):
        """
        logic:
            self._current_first_title_level > 0:     -> abstract
            not has_auto_num:                  -> abstract and appendix
        """
        if self._current_first_title_level > 0:
            if self._current_second_title_level == 0:
                self.save_section(self._grandpa_file)
            else:
                self.save_section(self._father_file)
        else:
            self.save_section(self._above_text)
        self._current_second_title_level = 0
        self._grandpa_file = self._father_file = ""

        if self.has_auto_num(para):
            self._current_first_title_level += 1
            num = str(self._current_first_title_level)
            self._last_level = level
            self._num_and_title = f"{num} {title}"
            self._first_router = self._num_and_title
            self.current_output_dir = os.path.join(self.output_doc_dir, self._first_router) + os.path.sep
            self._grandpa_file += f"# {self._num_and_title}\n"
        else:
            # 摘要和附录信息
            self._current_first_title_level = 0
            self._num_and_title = title.strip()
            self.current_output_dir = self.output_doc_dir
            self._last_level = 0
            self._first_router = ""
            self._grandpa_file += f"## {title.strip()}\n"
        
        return self._num_and_title.strip()
    
    def _parse_header_level_2(self, title: str):
        """
        self._current_first_title_level > 0:     -> not abstract
        """
        if self._current_first_title_level > 0:  # 排除摘要
            if self._current_second_title_level == 0:
                self.save_section(self._grandpa_file)
                self._grandpa_file = ""
            else:
                self.save_section(self._father_file)
                self._father_file = ""

        self._current_second_title_level += 1
        self._current_third_title_level = 0
        if self._last_level > 0:
            num = f"{self._current_first_title_level}.{self._current_second_title_level}"
            self._num_and_title = f"{num} {title}"
            self._second_router = self._num_and_title
            self.current_output_dir = os.path.join(
                self.output_doc_dir,
                self._first_router,
                self._second_router) + os.path.sep
            self._father_file += f"## {self._num_and_title}\n\n"
            return self._num_and_title.strip()
        else:
            self._father_file += f"## {title.strip()}\n\n"

    def _parse_header_level_3(self, title: str):
        """
        self._current_first_title_level > 0:     -> not abstract
        """
        if self._current_first_title_level > 0 and self._current_second_title_level > 0:
            self.save_section(self._father_file)
            self._father_file = ""

        self._current_third_title_level += 1
        num = f"{self._current_first_title_level}." \
              f"{self._current_second_title_level}." \
              f"{self._current_third_title_level}"
        if self._last_level > 0:
            self._num_and_title = f"{num} {title}"
            self._father_file += f"### {num} {title}\n\n"
            return self._num_and_title.strip()
        else:
            self._father_file += f"### {title}\n\n"

    def _parse_header(self, para):
        """
        processing headers of all level
        """
        level = self.get_heading_level(para)

        title_text = self.clean_title(para.text)
        if not title_text:
            return
        
        res = None
        if level == 0:
            self._parse_header_level_0(para)
        elif level == 1:
            res = self._parse_header_level_1(para, level=level, title=title_text)
        elif level == 2 and self._current_first_title_level > 0:
            res = self._parse_header_level_2(title=title_text)
        elif level == 3 and self._current_second_title_level > 0:
            res = self._parse_header_level_3(title=title_text)

        if res:
            self._frame.append(res)

    def parse_header_paragraph(self, para: Paragraph, level: int):
        """
        parse header paragraph
        """
        title_text = self.clean_title(para.text)
        if not title_text:
            return

        self._parse_header(para)
        self._num_and_title = re.sub(r'[\\/]', ' or ', self._num_and_title)
        self._num_and_title = re.sub(r' {2}', ' ', self._num_and_title)
        self._current_filename = f"{self._num_and_title}.md"
        logger.info(f"processing title: <level={level}>: {title_text}")

    def append_above_text(self, content: str, level: int):
        if level < 2:
            self._grandpa_file += content + "  \n\n"
        else:
            self._father_file += content + "  \n\n"

    def parse(self, export_path: str):
        """
        parse processing
        """
        logger.info(f"start to parsing docx file: {self.doc_name}")
        level = -1
        self._set_output_path(export_path)

        try:
            for para in self.iter_block_items(self.doc):
                self._block_count += 1
                if isinstance(para, Paragraph):
                    if self.is_heading(para):
                        self._heading_count += 1
                        level = self.get_heading_level(para)
                        self.parse_header_paragraph(para, level=level)
                    else:
                        paragraph_content = self.process_paragraph(para)
                        if paragraph_content.strip() and self._last_level > -1:
                            self.append_above_text(content=paragraph_content, level=level)
                        logger.info(f"processing normal paragraph: <level={level}>")
                elif isinstance(para, Table):
                    self._table_count += 1
                    logger.info(f"processing table: table<index={self._table_count}>")
                    table_info = self.process_table(para)
                    if table_info.strip() and self._last_level > -1:
                        self.append_above_text(content=table_info, level=level)

                # print(self._first_router, self._second_router, self.current_output_dir)
            self.save_last_section()
            self.save_frame()

            logger.info(f"file processing success: {self.doc_name}")
            logger.info(
                f"statistics: total_block={self._block_count}, "
                f"title_count={self._heading_count}, "
                f"table_count={self._table_count}, "
                f"image_count={self._image_count} "
            )
            if self.is_merge:
                self.merge_content()
        except Exception as e:
            logger.debug(f"DOCX process failed: {traceback.format_exc()}")
            logger.error(f"DOCX process failed: {e}")
            raise
    
    def save_last_section(self):
        """
        warning: should be called at the end of paragraph processing
        not strict logic!!!
        """
        if self._current_filename.strip():
            logger.info(f"tha last section: {self._current_filename}")
            self.save_section(self._above_text)
    
    def save_frame(self):
        """save frame file if self._frame is not empty
        """
        if not self._frame:
            return
        self._current_filename = f"{self.doc_name}_frame.md"
        md_text = "\n".join(self._frame)
        self.current_output_dir = self.output_doc_dir
        self.save_section(md_text)
        logger.info(f"save file frame: {self._current_filename}")

    def save_section(self, content: str):
        """save section as a file"""
        if not self._current_filename:
            logger.warning("skip save section: filename is None")
            return

        if not content.strip():
            logger.warning(f"skip empty section: {self._current_filename}")
            return
        
        logger.info(f"save section: {self._current_filename}")
        self._current_filename = self._current_filename.strip()
        try:
            file_content = self.escape_markdown_special_chars(content)
            if self._current_filename.startswith("0 "):
                self._current_filename = "COVERS_" + self._current_filename[2:]
            
            os.makedirs(self.current_output_dir, exist_ok=True)
            file_path = os.path.join(self.current_output_dir, self._current_filename)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(file_content + "  \n")
            logger.info(f"section saved: {file_path}")

        except Exception as e:
            logger.debug(f"save section failed: {self._current_filename}, error: {traceback.format_exc()}")
            logger.error(f"save section failed: {self._current_filename}, error: {e}")

    def escape_markdown_special_chars(self, text):
        """escape markdown special characters except code blocks and inline code"""
        try:
            code_blocks = []
            inline_codes = []
            code_block_pattern = re.compile(r'```.*?```', re.DOTALL)

            def code_block_replacer(match):
                code_blocks.append(match.group(0))
                return f"\x02CODE_BLOCK_{len(code_blocks) - 1}\x03"

            text = code_block_pattern.sub(code_block_replacer, text)
            inline_code_pattern = re.compile(r'`[^`]*`')

            def inline_code_replacer(match):
                inline_codes.append(match.group(0))
                return f"\x02INLINE_CODE_{len(inline_codes) - 1}\x03"

            text = inline_code_pattern.sub(inline_code_replacer, text)
            text = re.sub(r'(?<!\\)([<>{}])', r'\\\1', text)

            for i, code in enumerate(inline_codes):
                text = text.replace(f"\x02INLINE_CODE_{i}\x03", code)
            for j, block in enumerate(code_blocks):
                text = text.replace(f"\x02CODE_BLOCK_{j}\x03", block)
            return text
        except Exception as e:
            logger.debug(f"error occured while processing makrdown special chars: {traceback.format_exc()}")
            logger.error(f"error occured while processing makrdown special chars: {e}")
            return text

    def process_paragraph(self, para):
        if self.is_toc(para):
            return ""
        """处理普通段落"""
        try:
            text = para.text.strip()
            # 处理图片
            para_xml = para._p.xml
            if para._element.xpath('.//pic:pic') or ('w:object' in para_xml and 'v:imagedata r:id' in para_xml) or (
                    'w:pict' in para_xml and 'v:imagedata r:id' in para_xml):
                # if  para._element.xpath('.//pic:pic') or ( 'v:imagedata r:id' in para_xml) :
                logger.info("处理段落图片")
                md_img = self.process_image(para, False)
                if md_img:
                    text = text + "  " + md_img if text else md_img

            # 处理流程图

            return text
        except Exception as e:
            logger.error(f"处理段落失败: {e}")
            return ""

    def merge_content(self):
        """处理根目录下的所有一级标题文件夹"""
        logger.info(f"开始处理文档合并")
        # 遍历根目录下的所有子文件夹
        for folder_name in os.listdir(self.output_doc_dir):
            chapter_folder_path = os.path.join(self.output_doc_dir, folder_name)
            if not os.path.isdir(chapter_folder_path):
                continue
            logger.info(f"处理一级标题文件夹: {folder_name}")
            # 尝试合并文件夹内容
            self.merge_chapter_folder(chapter_folder_path)
        self.clean_marked_folders(self.output_doc_dir)

    def estimate_tokens(self, text: str, model_name: str = "cl100k_base") -> int:
        """使用tiktoken计算文本的token数"""
        try:
            os.environ["TIKTOKEN_CACHE_DIR"] = os.path.dirname(self._encoder_path)
            mergeable_ranks = tiktoken.load.load_tiktoken_bpe(self._encoder_path)
            token_encoder = tiktoken.Encoding(
                name=model_name,
                pat_str=r"""'(?i:[sdmt]|ll|ve|re)|[^\r\n\p{L}\p{N}]?+\p{L}+|\p{N}{1,3}| ?[^\s\p{L}\p{N}]++[\r\n]*|\s*[\r\n]|\s+(?!\S)|\s+'""",
                mergeable_ranks=mergeable_ranks,
                special_tokens={
                    "<|endoftext|>": 100257,
                    "<|fim_prefix|>": 100258,
                    "<|fim_middle|>": 100259,
                    "<|fim_suffix|>": 100260,
                    "<|endofprompt|>": 100276
                }
            )

            tokens = len(token_encoder.encode(text))
            logger.debug(f"估算token数量: 文本长度 {len(text)} 字符, 估算token数: {tokens}")
            return tokens

        except Exception as e:
            logger.error(f"使用本地编码器失败: {e}")

    def merge_chapter_folder(self, folder_path):
        """
        递归合并文件夹下的所有Markdown文档（包括子文件夹）
        """
        logger.info(f"开始合并文档,文件夹: {folder_path}")
        merged_result = "undo"

        # 递归处理所有子文件夹
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)
            if os.path.isdir(item_path):
                logger.info(f"发现子文件夹: {item_path}")
                results = self.merge_chapter_folder(item_path)
                if results == "stop":
                    merged_result = results

        if merged_result == "stop":
            return merged_result
            # 获取当前文件夹内所有md文件
        files = [f for f in os.listdir(folder_path) if f.endswith('.md')]
        # 按章节序号排序
        files.sort(key=lambda f: [int(num) for num in re.findall(r'\d+', f)])
        logger.info(f"找到 {len(files)} 个Markdown文件")

        # 空文件夹处理
        if files:
            merged_content = ""
            total_tokens = 0
            # 遍历文件夹中的所有Markdown文件
            for file in files:
                file_path = os.path.join(folder_path, file)
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                file_content = f"\n\n{content}\n\n"
                # 估算token数量
                file_tokens = self.estimate_tokens(file_content)
                logger.debug(f"文件 '{file}' 估算token: {file_tokens}")
                # 检查是否超过token限制
                if (total_tokens + file_tokens) > self.max_tokens:
                    logger.warning(f"合并内容超出token限制 ({total_tokens + file_tokens} > {self.max_tokens})，停止合并")
                    merged_result = "stop"
                merged_content += file_content
                total_tokens += file_tokens
                if len(files)>1 and merged_result == "stop":
                    return merged_result


            # 生成输出文件名并保存
            folder_name = os.path.basename(folder_path.rstrip('/\\'))
            output_file = os.path.join(os.path.dirname(folder_path), f"{folder_name}.md")
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(merged_content)
            logger.info(f"{folder_path} 合并完成，输出文件: {output_file}")
        else:
            logger.warning(f"文件夹为空: {folder_path}")

        # 复制图片，标记文件夹为删除
        self.copy_images(folder_path, os.path.dirname(folder_path))
        os.rename(folder_path, f"{folder_path}_delete")
        merged_result = "merged"
        return merged_result

    def copy_images(self, source_dir, output_dir):
        """
        复制图片文件到输出目录
        """
        logger.info(f"复制图片: {source_dir} -> {output_dir}")
        copied_files = []

        for file in os.listdir(source_dir):
            if file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
                src_path = os.path.join(source_dir, file)
                dest_path = os.path.join(output_dir, file)

                try:
                    shutil.copy2(src_path, dest_path)
                    copied_files.append(file)
                except Exception as e:
                    logger.error(f"复制图片失败: {src_path} -> {dest_path}, 错误: {e}")

        logger.info(f"复制了 {len(copied_files)} 张图片: {', '.join(copied_files)}")

    def clean_marked_folders(self, folder_path):
        """
        删除指定文件夹下所有以 '_delete' 结尾的文件夹（包括子文件夹中的）
        Args:
            folder_path: 要清理的根文件夹路径
            :param self:
        """
        logger.info(f"开始清理标记为删除的文件夹: {folder_path}")
        deleted_count = 0
        # 遍历文件夹中的所有项目
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)

            # 如果是文件夹
            if os.path.isdir(item_path):
                # 如果文件夹以 '_delete' 结尾，则删除
                if item.endswith('_delete'):
                    try:
                        shutil.rmtree(item_path)
                        deleted_count += 1
                    except Exception as e:
                        logger.error(f"删除文件夹失败 {item_path}: {e}")
                else:
                    # 递归检查子文件夹
                    self.clean_marked_folders(item_path)
        logger.info(f"清理完成，共删除 {deleted_count} 个文件夹")
        return deleted_count


if __name__ == "__main__":
    path = "/data/home/solgeo/temp/dv019_it_operation_guide_capdam_v2_5.docx"
    encoder_path = "/data/home/solgeo/models/tiktoken_cache/cl100k_base.tiktoken"
    output_path = "./temp/"

    logging.basicConfig(level=logging.INFO)

    wtm = WordSpliter(
        docx_object=path,
        max_tokens=5000,
        encoder_path=encoder_path
    )
    wtm.parse(export_path=output_path)
